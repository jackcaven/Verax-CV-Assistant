"""Unit tests for output builders."""

import tempfile
from pathlib import Path

import pytest
from docx import Document

from verax.builder.factory import BuilderFactory
from verax.builder.docx_builder import DocxBuilder
from verax.builder.html_builder import HtmlBuilder
from verax.builder.pdf_builder import detect_libreoffice
from verax.models.structured_cv import ContactInfo, CVEntry, CVSection, SectionType, StructuredCV
from verax.models.template_schema import SectionTemplate, TemplateSchema
from verax.template.fallback import get_fallback_template


@pytest.fixture
def contact_info():
    """Sample contact info."""
    return ContactInfo(
        name="John Doe",
        email="john@example.com",
        phone="+1 (555) 123-4567",
        location="New York, NY",
        website="https://johndoe.com",
        linkedin="https://linkedin.com/in/johndoe",
    )


@pytest.fixture
def sample_sections():
    """Sample CV sections."""
    experience_entries = [
        CVEntry(
            title="Software Engineer",
            subtitle="Tech Corp",
            dates="Jan 2020 - Present",
            description="• Led development of new features\n• Improved system performance by 40%",
        ),
        CVEntry(
            title="Junior Developer",
            subtitle="StartUp Inc",
            dates="Jun 2018 - Dec 2019",
            description="• Contributed to backend systems\n• Fixed critical bugs",
        ),
    ]

    education_entries = [
        CVEntry(
            title="Bachelor of Science",
            subtitle="University of Technology",
            dates="2014 - 2018",
            description="Major: Computer Science\nGPA: 3.8/4.0",
        ),
    ]

    skills_entries = [
        CVEntry(title="Python", subtitle="Expert"),
        CVEntry(title="JavaScript", subtitle="Advanced"),
        CVEntry(title="Docker", subtitle="Intermediate"),
    ]

    return [
        CVSection(
            title="Experience",
            section_type=SectionType.EXPERIENCE,
            entries=experience_entries,
        ),
        CVSection(
            title="Education",
            section_type=SectionType.EDUCATION,
            entries=education_entries,
        ),
        CVSection(
            title="Skills",
            section_type=SectionType.SKILLS,
            entries=skills_entries,
        ),
    ]


@pytest.fixture
def structured_cv(contact_info, sample_sections):
    """Sample structured CV."""
    return StructuredCV(
        contact_info=contact_info,
        sections=sample_sections,
        source_filename="sample.pdf",
    )


@pytest.fixture
def template_schema():
    """Sample template schema."""
    return TemplateSchema(
        sections=[
            SectionTemplate(
                title="Experience",
                heading_style="Heading 1",
                font_name="Calibri",
                font_size=14,
                order_index=0,
            ),
            SectionTemplate(
                title="Education",
                heading_style="Heading 1",
                font_name="Calibri",
                font_size=14,
                order_index=1,
            ),
            SectionTemplate(
                title="Skills",
                heading_style="Heading 1",
                font_name="Calibri",
                font_size=14,
                order_index=2,
            ),
        ],
        raw_docx_path=str(get_fallback_template()),
        source_format="docx",
        contact_block_style="Normal",
    )


class TestBuilderFactory:
    """Test BuilderFactory."""

    def test_create_docx_builder(self):
        """Test creating DOCX builder."""
        builder = BuilderFactory.create("docx")
        assert builder.name == "docx"

    def test_create_html_builder(self):
        """Test creating HTML builder."""
        builder = BuilderFactory.create("html")
        assert builder.name == "html"

    def test_create_pdf_builder(self):
        """Test creating PDF builder."""
        builder = BuilderFactory.create("pdf")
        assert builder.name == "pdf"

    def test_unsupported_format(self):
        """Test error on unsupported format."""
        with pytest.raises(ValueError, match="Unsupported output format"):
            BuilderFactory.create("unknown")

    def test_get_supported_formats(self):
        """Test getting supported formats."""
        formats = BuilderFactory.get_supported_formats()
        assert "docx" in formats
        assert "html" in formats
        assert "pdf" in formats

    def test_is_available_docx(self):
        """Test DOCX always available."""
        assert BuilderFactory.is_available("docx")

    def test_is_available_html(self):
        """Test HTML always available."""
        assert BuilderFactory.is_available("html")

    def test_is_available_pdf(self):
        """Test PDF availability (depends on LibreOffice)."""
        # PDF available only if LibreOffice installed
        available = BuilderFactory.is_available("pdf")
        if detect_libreoffice():
            assert available
        # If LibreOffice not installed, that's ok for test


class TestDocxBuilder:
    """Test DOCX builder."""

    def test_build_creates_file(self, structured_cv, template_schema):
        """Test that build() creates DOCX file."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.docx"
            builder = DocxBuilder()
            builder.build(structured_cv, template_schema, output_path)

            assert output_path.exists()
            assert output_path.suffix == ".docx"

    def test_built_docx_is_valid(self, structured_cv, template_schema):
        """Test that built DOCX is valid and readable."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.docx"
            builder = DocxBuilder()
            builder.build(structured_cv, template_schema, output_path)

            # Verify it's a valid DOCX by loading it
            doc = Document(output_path)
            assert len(doc.paragraphs) > 0

    def test_contact_info_in_output(self, structured_cv, template_schema):
        """Test that contact info appears in output."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.docx"
            builder = DocxBuilder()
            builder.build(structured_cv, template_schema, output_path)

            doc = Document(output_path)
            text = "\n".join(p.text for p in doc.paragraphs)

            assert structured_cv.contact_info.name in text
            assert structured_cv.contact_info.email in text

    def test_sections_in_output(self, structured_cv, template_schema):
        """Test that CV sections appear in output."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.docx"
            builder = DocxBuilder()
            builder.build(structured_cv, template_schema, output_path)

            doc = Document(output_path)
            text = "\n".join(p.text for p in doc.paragraphs)

            assert "Experience" in text
            assert "Education" in text
            assert "Skills" in text

    def test_entries_in_output(self, structured_cv, template_schema):
        """Test that CV entries appear in output."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.docx"
            builder = DocxBuilder()
            builder.build(structured_cv, template_schema, output_path)

            doc = Document(output_path)
            text = "\n".join(p.text for p in doc.paragraphs)

            # Check for experience entry
            assert "Software Engineer" in text
            assert "Tech Corp" in text

    def test_missing_template_uses_fallback(self, structured_cv):
        """Test that missing template path falls back to default."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.docx"

            # Use template with no raw_docx_path
            template = TemplateSchema(
                sections=[
                    SectionTemplate(
                        title="Experience",
                        heading_style="Heading 1",
                        order_index=0,
                    )
                ],
                raw_docx_path=None,
                source_format="docx",
            )

            builder = DocxBuilder()
            builder.build(structured_cv, template, output_path)

            assert output_path.exists()

    def test_invalid_template_path_raises(self, structured_cv):
        """Test error when template path invalid."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.docx"

            template = TemplateSchema(
                sections=[],
                raw_docx_path="/nonexistent/path/template.docx",
                source_format="docx",
            )

            builder = DocxBuilder()
            with pytest.raises(ValueError, match="Template not found"):
                builder.build(structured_cv, template, output_path)


class TestHtmlBuilder:
    """Test HTML builder."""

    def test_build_creates_file(self, structured_cv, template_schema):
        """Test that build() creates HTML file."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.html"
            builder = HtmlBuilder()
            builder.build(structured_cv, template_schema, output_path)

            assert output_path.exists()
            assert output_path.suffix == ".html"

    def test_html_is_valid(self, structured_cv, template_schema):
        """Test that HTML is valid."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.html"
            builder = HtmlBuilder()
            builder.build(structured_cv, template_schema, output_path)

            html_content = output_path.read_text()
            assert "<!DOCTYPE html>" in html_content
            assert "</html>" in html_content

    def test_contact_info_in_html(self, structured_cv, template_schema):
        """Test that contact info appears in HTML."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.html"
            builder = HtmlBuilder()
            builder.build(structured_cv, template_schema, output_path)

            html_content = output_path.read_text()
            assert structured_cv.contact_info.name in html_content
            assert structured_cv.contact_info.email in html_content

    def test_sections_in_html(self, structured_cv, template_schema):
        """Test that sections appear in HTML."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.html"
            builder = HtmlBuilder()
            builder.build(structured_cv, template_schema, output_path)

            html_content = output_path.read_text()
            assert "Experience" in html_content
            assert "Education" in html_content

    def test_entries_in_html(self, structured_cv, template_schema):
        """Test that entries appear in HTML."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.html"
            builder = HtmlBuilder()
            builder.build(structured_cv, template_schema, output_path)

            html_content = output_path.read_text()
            assert "Software Engineer" in html_content
            assert "Tech Corp" in html_content

    def test_html_has_css(self, structured_cv, template_schema):
        """Test that HTML includes CSS styling."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "output.html"
            builder = HtmlBuilder()
            builder.build(structured_cv, template_schema, output_path)

            html_content = output_path.read_text()
            assert "<style>" in html_content
            assert "body {" in html_content

    def test_creates_parent_directories(self, structured_cv, template_schema):
        """Test that build() creates parent directories if needed."""
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = Path(tmpdir) / "subdir" / "deep" / "output.html"
            builder = HtmlBuilder()
            builder.build(structured_cv, template_schema, output_path)

            assert output_path.exists()
