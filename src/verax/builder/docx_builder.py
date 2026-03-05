"""Clone-and-fill DOCX builder."""

import shutil
from pathlib import Path

from docx import Document

from verax.models.structured_cv import CVEntry, CVSection, StructuredCV
from verax.models.template_schema import TemplateSchema
from verax.template.fallback import get_fallback_template
from verax.utils import get_logger

logger = get_logger(__name__)


class DocxBuilder:
    """Build DOCX output by cloning template and filling with CV data."""

    name = "docx"

    def build(
        self,
        cv: StructuredCV,
        template_schema: TemplateSchema,
        output_path: Path,
    ) -> None:
        """Build DOCX by cloning template and filling with CV data.

        Process:
        1. Copy template DOCX to output path
        2. Remove non-heading paragraphs (keep section headings as references)
        3. Rebuild from StructuredCV using heading styles from template
        4. Preserve all formatting from original template

        Args:
            cv: Structured CV to write
            template_schema: Template structure with section headings
            output_path: Where to save the DOCX file

        Raises:
            ValueError: If template path is invalid or CV data is incomplete
            IOError: If file operations fail
        """
        logger.info(f"Building DOCX for {cv.source_filename} -> {output_path}")

        # Get template path (from extraction or fallback)
        template_path = self._get_template_path(template_schema)
        if not template_path.exists():
            raise ValueError(f"Template not found: {template_path}")

        # Copy template to output path
        try:
            shutil.copy2(template_path, output_path)
            logger.debug(f"Copied template to {output_path}")
        except Exception as e:
            raise IOError(f"Failed to copy template: {e}")

        # Load copied document
        try:
            doc = Document(output_path)
        except Exception as e:
            raise ValueError(f"Failed to load template DOCX: {e}")

        # Clear body paragraphs (keep structure/headings as style references)
        self._clear_body(doc)

        # Rebuild document
        try:
            self._add_contact_block(doc, cv.contact_info, template_schema)
            self._add_sections(doc, cv.sections, template_schema)
        except Exception as e:
            logger.error(f"Failed to rebuild document: {e}")
            raise

        # Save
        try:
            doc.save(output_path)
            logger.info(f"DOCX saved: {output_path}")
        except Exception as e:
            raise IOError(f"Failed to save DOCX: {e}")

    @staticmethod
    def _get_template_path(template_schema: TemplateSchema) -> Path:
        """Get template DOCX path, fallback if not provided.

        Args:
            template_schema: Template schema with optional raw_docx_path

        Returns:
            Path to template DOCX file
        """
        if template_schema.raw_docx_path:
            return Path(template_schema.raw_docx_path)
        else:
            return get_fallback_template()

    @staticmethod
    def _clear_body(doc: Document) -> None:
        """Remove all body paragraphs except those with heading styles.

        Preserves heading styles as structural references for later insertion.

        Args:
            doc: Document object to clear
        """
        paras_to_remove = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            # Keep heading paragraphs as style references; remove others
            if not style_name.startswith("Heading"):
                paras_to_remove.append(para)

        # Remove marked paragraphs
        for para in paras_to_remove:
            p = para._element
            p.getparent().remove(p)

        logger.debug(f"Cleared {len(paras_to_remove)} non-heading paragraphs")

    @staticmethod
    def _add_contact_block(doc: Document, contact_info, template_schema: TemplateSchema) -> None:
        """Add contact information block at the beginning of document.

        Args:
            doc: Document to add contact block to
            contact_info: ContactInfo object
            template_schema: Template schema for style hints
        """
        # Insert contact block at the start
        contact_parts = []

        if contact_info.name:
            contact_parts.append(contact_info.name)

        if contact_info.email:
            contact_parts.append(contact_info.email)

        if contact_info.phone:
            contact_parts.append(contact_info.phone)

        if contact_info.location:
            contact_parts.append(contact_info.location)

        if contact_info.website:
            contact_parts.append(contact_info.website)

        if contact_info.linkedin:
            contact_parts.append(contact_info.linkedin)

        if contact_parts:
            # Add name as first paragraph
            name_para = doc.add_paragraph(contact_info.name or "")
            name_para.style = template_schema.contact_block_style or "Normal"

            # Add other info as separate lines or combined
            contact_line = " | ".join(contact_parts[1:]) if len(contact_parts) > 1 else ""
            if contact_line:
                info_para = doc.add_paragraph(contact_line)
                info_para.style = "Normal"

        logger.debug(f"Added contact block for {contact_info.name}")

    @staticmethod
    def _add_sections(
        doc: Document,
        sections: list[CVSection],
        template_schema: TemplateSchema,
    ) -> None:
        """Add CV sections to document.

        For each template section, find matching CV section and add entries.
        If CV missing a template section, add empty placeholder.

        Args:
            doc: Document to add sections to
            sections: CV sections to add
            template_schema: Template structure with section order
        """
        for template_section in template_schema.sections:
            # Find matching CV section by title
            matching_cv_section = next(
                (s for s in sections if s.title.lower() == template_section.title.lower()),
                None,
            )

            if matching_cv_section:
                DocxBuilder._add_section(doc, template_section, matching_cv_section)
            else:
                DocxBuilder._add_placeholder_section(doc, template_section)

        logger.debug(f"Added {len(template_schema.sections)} sections")

    @staticmethod
    def _add_section(doc: Document, template_section, cv_section: CVSection) -> None:
        """Add a single section with entries.

        Args:
            doc: Document to add to
            template_section: SectionTemplate with style info
            cv_section: CVSection with entries
        """
        # Add section heading (reuse style from template)
        heading_para = doc.add_paragraph(template_section.title)
        heading_para.style = template_section.heading_style or "Heading 1"

        # Add entries
        for entry in cv_section.entries:
            DocxBuilder._add_entry(doc, entry)

        logger.debug(f"Added section: {cv_section.title} with {len(cv_section.entries)} entries")

    @staticmethod
    def _add_entry(doc: Document, entry: CVEntry) -> None:
        """Add a single CV entry (title, subtitle, dates, description).

        Args:
            doc: Document to add to
            entry: CVEntry to add
        """
        # Title + subtitle on same line
        title_text = entry.title
        if entry.subtitle:
            title_text = f"{entry.title} — {entry.subtitle}"

        title_para = doc.add_paragraph(title_text)
        title_para.style = "Normal"

        # Make title bold
        for run in title_para.runs:
            run.bold = True

        # Dates (if present)
        if entry.dates:
            dates_para = doc.add_paragraph(entry.dates)
            dates_para.style = "Normal"

        # Description (bullet points)
        if entry.description:
            for line in entry.description.split("\n"):
                line = line.strip()
                if line:
                    # Remove existing bullet symbols if present
                    if line.startswith("•"):
                        line = line[1:].strip()
                    if line.startswith("-"):
                        line = line[1:].strip()

                    doc.add_paragraph(line, style="List Bullet")

    @staticmethod
    def _add_placeholder_section(doc: Document, template_section) -> None:
        """Add empty placeholder for missing section.

        Args:
            doc: Document to add to
            template_section: SectionTemplate to create placeholder for
        """
        heading_para = doc.add_paragraph(template_section.title)
        heading_para.style = template_section.heading_style or "Heading 1"

        placeholder_para = doc.add_paragraph("[Content for this section would go here]")
        placeholder_para.style = "Normal"
        for run in placeholder_para.runs:
            run.italic = True

        logger.debug(f"Added placeholder for section: {template_section.title}")
