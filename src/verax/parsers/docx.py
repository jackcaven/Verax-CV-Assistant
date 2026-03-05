"""Parser for DOCX documents."""

from pathlib import Path

from docx import Document

from verax.parsers.models import RawDocument
from verax.utils import get_logger

logger = get_logger(__name__)


class DoxcParser:
    """Parser for DOCX documents using python-docx."""

    def parse(self, file_path: Path) -> RawDocument:
        """Parse DOCX file and extract plain text.

        Extracts text from all paragraphs and tables in order.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            RawDocument with extracted text.

        Raises:
            FileNotFoundError: If file does not exist.
            ValueError: If file is not a valid DOCX.
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {e}")

        text_parts: list[str] = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        full_text = "\n".join(text_parts)
        logger.debug(f"Parsed DOCX: {file_path.name}, {len(full_text)} chars")

        return RawDocument(
            text=full_text,
            source_filename=file_path.name,
            source_format="docx",
        )
